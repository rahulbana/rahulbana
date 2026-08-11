"""Render a ReviewReport as a Microsoft Word (.docx) document."""

from __future__ import annotations

from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor

from ..models import ReviewReport, Severity

_SEVERITY_LABEL = {
    Severity.CRITICAL: "CRITICAL",
    Severity.HIGH: "HIGH",
    Severity.MEDIUM: "MEDIUM",
    Severity.LOW: "LOW",
    Severity.INFO: "INFO",
}

_SEVERITY_COLOR = {
    Severity.CRITICAL: RGBColor(0xC0, 0x00, 0x00),
    Severity.HIGH: RGBColor(0xE0, 0x5A, 0x00),
    Severity.MEDIUM: RGBColor(0xB8, 0x8A, 0x00),
    Severity.LOW: RGBColor(0x00, 0x70, 0xA0),
    Severity.INFO: RGBColor(0x70, 0x70, 0x70),
}


def write_docx(report: ReviewReport, path: str) -> None:
    """Build and save a .docx report to ``path``."""
    document = Document()

    document.add_heading("AI Code Review", level=0)

    meta = document.add_paragraph()
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta.add_run("Path: ").bold = True
    meta.add_run(f"{report.root}\n")
    meta.add_run("Language: ").bold = True
    meta.add_run(f"{report.language}\n")
    meta.add_run("Model: ").bold = True
    meta.add_run(f"{report.model}\n")
    meta.add_run("Files reviewed: ").bold = True
    meta.add_run(f"{len(report.files)}\n")
    meta.add_run("Generated: ").bold = True
    meta.add_run(generated)

    # Summary table.
    document.add_heading("Summary", level=1)
    counts = report.severity_counts()
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Count"
    for sev in Severity:
        row = table.add_row().cells
        run = row[0].paragraphs[0].add_run(_SEVERITY_LABEL[sev])
        run.font.color.rgb = _SEVERITY_COLOR[sev]
        run.bold = True
        row[1].text = str(counts[sev])
    total_row = table.add_row().cells
    total_run = total_row[0].paragraphs[0].add_run("TOTAL")
    total_run.bold = True
    total_count = total_row[1].paragraphs[0].add_run(
        str(len(report.all_findings))
    )
    total_count.bold = True

    # Per-file detail.
    document.add_heading("Findings", level=1)
    for file_review in report.files:
        document.add_heading(file_review.file_path, level=2)
        findings = file_review.findings
        errors = [r for r in file_review.results if r.error]

        if not findings and not errors:
            document.add_paragraph("No issues found.")
            continue

        for result in file_review.results:
            if result.error:
                para = document.add_paragraph()
                run = para.add_run(
                    f"{result.category} review failed: {result.error}"
                )
                run.italic = True
                continue
            if not result.findings:
                continue

            document.add_heading(result.category, level=3)
            for finding in sorted(
                result.findings, key=lambda f: f.severity.rank
            ):
                _add_finding(document, finding)

    document.save(path)


def _add_finding(document: "Document", finding) -> None:
    para = document.add_paragraph(style="List Bullet")
    badge = para.add_run(f"[{_SEVERITY_LABEL[finding.severity]}] ")
    badge.bold = True
    badge.font.color.rgb = _SEVERITY_COLOR[finding.severity]

    title = para.add_run(finding.title)
    title.bold = True
    if finding.line:
        para.add_run(f" (line {finding.line})").italic = True

    if finding.description:
        document.add_paragraph(finding.description)
    if finding.recommendation:
        rec = document.add_paragraph()
        label = rec.add_run("Recommendation: ")
        label.bold = True
        rec.add_run(finding.recommendation)
