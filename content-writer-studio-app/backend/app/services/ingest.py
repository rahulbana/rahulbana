"""Extract plain text from uploaded files and links for style references."""
from __future__ import annotations

import io
import re

from fastapi import HTTPException, status

# Cap stored text so a huge page/file can't bloat the DB or a prompt.
MAX_CHARS = 200_000

ALLOWED_FILE_EXTENSIONS = (".txt", ".md", ".docx")


def _clean(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_CHARS]


# --- Files ----------------------------------------------------------

def extract_from_file(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".docx"):
        return _clean(_extract_docx(data))
    if name.endswith((".txt", ".md")):
        return _clean(data.decode("utf-8", errors="ignore"))
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Unsupported file type. Upload a .txt, .md or .docx file.",
    )


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Could not read .docx file: {exc}",
        ) from exc


# --- Links ----------------------------------------------------------

_SCRIPT_STYLE = re.compile(r"<(script|style|noscript)[^>]*>.*?</\1>", re.I | re.S)
_TAGS = re.compile(r"<[^>]+>")
_TITLE = re.compile(r"<title[^>]*>(.*?)</title>", re.I | re.S)


def extract_from_url(url: str) -> tuple[str, str]:
    """Fetch a URL and return (title, cleaned_text)."""
    if not re.match(r"^https?://", url, re.I):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Please provide a valid http(s) URL.",
        )
    import httpx

    try:
        resp = httpx.get(
            url,
            timeout=20,
            follow_redirects=True,
            headers={"User-Agent": "AIContentWriter/1.0 (+style-reference)"},
        )
        resp.raise_for_status()
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Could not fetch the URL: {exc}",
        ) from exc

    html = resp.text
    title_match = _TITLE.search(html)
    title = _clean_inline(title_match.group(1)) if title_match else url

    body = _SCRIPT_STYLE.sub(" ", html)
    body = _TAGS.sub(" ", body)
    body = _unescape(body)
    text = _clean(body)
    if not text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No readable text found at that URL.",
        )
    return title or url, text


def _clean_inline(text: str) -> str:
    return re.sub(r"\s+", " ", _unescape(text)).strip()


def _unescape(text: str) -> str:
    import html as html_lib

    return html_lib.unescape(text)
