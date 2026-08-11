"""Export an article to PDF, DOCX or Markdown (returns raw bytes)."""
from __future__ import annotations

import io
import re

from app.models.article import Article
from app.services.images import resolve_media_path


def _source_label(src: dict) -> str:
    title = src.get("title") or "Untitled source"
    url = src.get("url")
    stype = src.get("type")
    label = f"{title} ({url})" if url else title
    return f"[{stype}] {label}" if stype else label


def to_markdown(article: Article) -> bytes:
    kw = ", ".join(article.keywords or [])
    tags = ", ".join(article.tags or [])
    banner = f"![{article.title}]({article.banner_image})\n\n" if article.banner_image else ""
    front = (
        banner
        + f"# {article.title}\n\n"
        f"> {article.summary}\n\n"
        f"**Keywords:** {kw}\n\n"
        f"**Tags:** {tags}\n\n"
        f"**Sentiment:** {article.sentiment}\n\n"
        "---\n\n"
    )
    body = article.body or ""
    if article.sources:
        lines = "\n".join(
            f"- {_source_label(s)}" for s in article.sources if isinstance(s, dict)
        )
        body += f"\n\n## Sources\n\n{lines}\n"
    return (front + body).encode("utf-8")


def _strip_md(text: str) -> str:
    """Very light inline Markdown -> plain text for PDF/DOCX bodies."""
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`{1,3}(.*?)`{1,3}", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text.strip()


# Parsed block kinds: ("h", text, level) | ("p", text, 0) |
# ("ul", text, 0) | ("ol", text, 0) | ("quote", text, 0) | ("code", text, 0)
def _parse_blocks(body: str) -> list[tuple[str, str, int]]:
    """Parse Markdown line-by-line into blocks, preserving all content.

    Unlike naive blank-line splitting, this never drops the prose that
    follows a heading, and it keeps every list item and paragraph.
    """
    blocks: list[tuple[str, str, int]] = []
    para: list[str] = []
    code_lines: list[str] = []
    in_code = False

    def flush_para() -> None:
        if para:
            blocks.append(("p", _strip_md(" ".join(para)), 0))
            para.clear()

    def flush_code() -> None:
        if code_lines:
            blocks.append(("code", "\n".join(code_lines), 0))
            code_lines.clear()

    for raw in (body or "").split("\n"):
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_para()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not stripped:
            flush_para()
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading:
            flush_para()
            blocks.append(("h", _strip_md(heading.group(2)), min(len(heading.group(1)), 4)))
            continue

        bullet = re.match(r"^[-*+]\s+(.*)", stripped)
        if bullet:
            flush_para()
            blocks.append(("ul", _strip_md(bullet.group(1)), 0))
            continue

        numbered = re.match(r"^\d+[.)]\s+(.*)", stripped)
        if numbered:
            flush_para()
            blocks.append(("ol", _strip_md(numbered.group(1)), 0))
            continue

        if stripped.startswith(">"):
            flush_para()
            blocks.append(("quote", _strip_md(stripped.lstrip(">").strip()), 0))
            continue

        if re.match(r"^([-*_])\1{2,}$", stripped):  # horizontal rule
            flush_para()
            continue

        para.append(stripped)

    flush_para()
    flush_code()
    return blocks


def to_docx(article: Article) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    banner_path = resolve_media_path(article.banner_image)
    if banner_path:
        from docx.shared import Inches

        try:
            doc.add_picture(banner_path, width=Inches(6.0))
        except Exception:  # pragma: no cover - never fail export on image
            pass
    doc.add_heading(article.title or "Untitled", level=0)
    if article.summary:
        p = doc.add_paragraph(article.summary)
        p.runs[0].italic = True

    for kind, text, level in _parse_blocks(article.body):
        if not text:
            continue
        if kind == "h":
            doc.add_heading(text, level=level)
        elif kind == "ul":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "ol":
            doc.add_paragraph(text, style="List Number")
        elif kind == "quote":
            doc.add_paragraph(text, style="Intense Quote")
        elif kind == "code":
            p = doc.add_paragraph(text)
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(9)
        else:
            doc.add_paragraph(text)

    if article.keywords:
        doc.add_heading("Keywords", level=2)
        doc.add_paragraph(", ".join(article.keywords)).runs[0].font.size = Pt(10)

    if article.sources:
        doc.add_heading("Sources", level=2)
        for s in article.sources:
            if isinstance(s, dict):
                doc.add_paragraph(_source_label(s), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(article: Article) -> bytes:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, title=article.title or "Untitled")
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=11, leading=16, alignment=TA_LEFT
    )

    flow: list = []
    banner_path = resolve_media_path(article.banner_image)
    if banner_path:
        try:
            img = Image(banner_path, width=468, height=312)  # 3:2 within letter margins
            img.hAlign = "CENTER"
            flow.append(img)
            flow.append(Spacer(1, 12))
        except Exception:  # pragma: no cover
            pass
    flow.append(Paragraph(_escape(article.title or "Untitled"), styles["Title"]))
    if article.summary:
        flow.append(Paragraph(f"<i>{_escape(article.summary)}</i>", styles["Italic"]))
    flow.append(Spacer(1, 12))

    bullet_style = ParagraphStyle(
        "Bullet", parent=body_style, leftIndent=16, bulletIndent=4, spaceBefore=2
    )
    quote_style = ParagraphStyle(
        "Quote", parent=body_style, leftIndent=16, textColor="#555555", fontName="Helvetica-Oblique"
    )
    code_style = ParagraphStyle(
        "Code", parent=styles["Code"], fontSize=9, leading=12, backColor="#f4f4f4"
    )

    for kind, text, level in _parse_blocks(article.body):
        if not text:
            continue
        if kind == "h":
            flow.append(Paragraph(_escape(text), styles[f"Heading{level}"]))
        elif kind == "ul":
            flow.append(Paragraph(_escape(text), bullet_style, bulletText="•"))
        elif kind == "ol":
            flow.append(Paragraph(_escape(text), bullet_style, bulletText="–"))
        elif kind == "quote":
            flow.append(Paragraph(_escape(text), quote_style))
        elif kind == "code":
            flow.append(Paragraph(_escape(text).replace("\n", "<br/>"), code_style))
        else:
            flow.append(Paragraph(_escape(text), body_style))
        flow.append(Spacer(1, 6))

    if article.sources:
        flow.append(Spacer(1, 12))
        flow.append(Paragraph("Sources", styles["Heading2"]))
        for s in article.sources:
            if isinstance(s, dict):
                flow.append(Paragraph(f"• {_escape(_source_label(s))}", body_style))

    doc.build(flow)
    return buf.getvalue()


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )


EXPORTERS = {
    "md": (to_markdown, "text/markdown", "md"),
    "markdown": (to_markdown, "text/markdown", "md"),
    "docx": (
        to_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "pdf": (to_pdf, "application/pdf", "pdf"),
}
