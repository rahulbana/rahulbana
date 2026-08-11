"""Render a `ResearchReport` to a Word (.docx) document."""

from __future__ import annotations

from pathlib import Path

from app.models.report import ResearchReport


def export_docx(report: ResearchReport, path: str | Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(report.title, level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Topic: {report.topic}").italic = True
    meta.add_run(f"   ·   Confidence: {report.confidence:.0%}").italic = True

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.executive_summary)

    if report.key_findings:
        doc.add_heading("Key Findings", level=1)
        for kf in report.key_findings:
            doc.add_paragraph(kf, style="List Bullet")

    for sec in report.sections:
        doc.add_heading(sec.heading, level=1)
        doc.add_paragraph(sec.body)
        for tbl in sec.tables:
            doc.add_heading(tbl.title, level=2)
            table = doc.add_table(rows=1, cols=len(tbl.columns))
            table.style = "Light Grid Accent 1"
            for j, col in enumerate(tbl.columns):
                cell = table.rows[0].cells[j]
                cell.text = col
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            for row in tbl.rows:
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)

    if report.open_questions:
        doc.add_heading("Open Questions & Gaps", level=1)
        for q in report.open_questions:
            doc.add_paragraph(q, style="List Bullet")

    if report.citations:
        doc.add_heading("References", level=1)
        for c in sorted(report.citations, key=lambda x: x.number):
            doc.add_paragraph(f"{c.number}. {c.title} — {c.url}")

    path = Path(path)
    doc.save(str(path))
    return path
