"""Loaders that normalise PDF, Word, Markdown and plain-text into a string.

The rest of the pipeline only ever sees text, so adding a new source format
means adding one function here and a branch in :func:`load_source`.
"""

from __future__ import annotations

from pathlib import Path

_MAX_CHARS = 60_000  # Keep prompts within a sane token budget.


def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) <= _MAX_CHARS:
        return text
    return text[:_MAX_CHARS] + "\n\n[... source truncated for length ...]"


def load_text(raw: str) -> str:
    """Pass-through for pasted / piped text."""
    return _truncate(raw)


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader  # local import keeps optional deps lazy

    reader = PdfReader(str(path))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(parts)


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Include table cell text too — proposals often hide content in tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_plain,
    ".md": _load_plain,
    ".markdown": _load_plain,
}


def load_source(path: str | Path) -> str:
    """Load a file by extension and return normalised text.

    Raises a clear error for unsupported formats (e.g. legacy ``.doc``).
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Source file not found: {p}")

    loader = _LOADERS.get(p.suffix.lower())
    if loader is None:
        supported = ", ".join(sorted(_LOADERS))
        raise ValueError(
            f"Unsupported file type '{p.suffix}'. Supported: {supported}. "
            "For .doc, please re-save as .docx or paste the text directly."
        )

    return _truncate(loader(p))
